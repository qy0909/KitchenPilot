import io
import os
from pathlib import Path

import streamlit as st
from jamaibase import JamAI, types as t, protocol as p
import json

st.set_page_config(page_title="Smart F&B Compliance Auditor", layout="wide")
st.title("Smart Auditor")
st.caption("Upload your SOP and optionally an official compliance doc. The system summarizes gaps.")


# CONFIG
PAT = st.secrets["PAT"]
PROJECT_ID = st.secrets["PROJECT_ID"]
AUDIT_TABLE_ID = "Audit"               # Your Action Table ID
INPUT_COL_USER = "user_sop"          # file col
OUTPUT_COL_SUMMARY = "gap_summary"          # text col

# Optional second reference doc
INPUT_COL_REF = "ref_compliance"       # file col

jamai = JamAI(token=PAT, project_id=PROJECT_ID)

# Helper: robust text extraction
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    import docx
except Exception:
    docx = None

def _extract_text(uploaded_file) -> str:
    if not uploaded_file:
        return ""
    try:
        uploaded_file.seek(0)
    except Exception:
        pass
    name = (getattr(uploaded_file, "name", "uploaded") or "uploaded").lower()
    ext = os.path.splitext(name)[1]

    def _decode_bytes(b: bytes) -> str:
        try:
            return b.decode("utf-8")
        except Exception:
            try:
                return b.decode("latin-1")
            except Exception:
                return "(Unable to decode file)"

    if ext == ".pdf" and PdfReader is not None:
        try:
            uploaded_file.seek(0)
            reader = PdfReader(uploaded_file)
            parts = []
            for page in reader.pages:
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                if txt:
                    parts.append(txt)
            if parts:
                return "\n\n".join(parts)[:20000]
        except Exception:
            pass

    if ext == ".docx" and docx is not None:
        try:
            uploaded_file.seek(0)
            d = docx.Document(uploaded_file)
            paras = [p.text for p in d.paragraphs if p.text]
            if paras:
                return "\n".join(paras)[:20000]
        except Exception:
            pass

    if ext == ".json":
        try:
            uploaded_file.seek(0)
            raw = uploaded_file.read()
            data = json.loads(_decode_bytes(raw))
            return json.dumps(data, indent=2)[:20000]
        except Exception:
            pass

    try:
        uploaded_file.seek(0)
        return _decode_bytes(uploaded_file.read())[:20000]
    except Exception:
        return "(Failed to read file)"


# Sidebar
with st.sidebar:
    st.markdown(
        """
        <div style='text-align: center; color: grey; font-size: 12px;'>
            Kitchen Pilot © 2025 | By CodeFest2025 Hackathon Team AAA
        </div>
        """,
        unsafe_allow_html=True
    )


# Main UI
left, right = st.columns(2)
with left:
    st.subheader("1) Upload your SOP (required)")
    user_file = st.file_uploader(
        "Your SOP document (txt/md/pdf/docx)",
        type=["txt", "md", "pdf", "docx", "csv", "log", "json"],
        key="user_sop"
    )

with right:
    st.subheader("2) Upload official compliance (optional)")
    ref_file = st.file_uploader(
        "Official compliance doc (txt/md/pdf/docx)",
        type=["txt", "md", "pdf", "docx", "csv", "log", "json"],
        key="ref_doc"
    )

analyze = st.button("Analyze Gaps", type="primary")
result_box = st.empty()


# Run Audit
if analyze:

    if not user_file:
        st.warning("Please upload your SOP document.")
        st.stop()

    try:
        with st.spinner("Sending to JamAI Base…"):
            user_text = _extract_text(user_file)
            ref_text = _extract_text(ref_file) if ref_file else ""

            row_data = {INPUT_COL_USER: user_text}
            if ref_text:
                row_data[INPUT_COL_REF] = ref_text

            req = t.MultiRowAddRequest(table_id=AUDIT_TABLE_ID, data=[row_data], stream=False)
            res = jamai.table.add_table_rows(t.TableType.ACTION, req)
    except Exception as e:
        st.error(f"Audit failed: {e}")
        st.stop()

    # Read model output
    row0 = res.rows[0] if hasattr(res, "rows") and res.rows else None
    cols = row0.columns if row0 else {}

    summary_cell = cols.get(OUTPUT_COL_SUMMARY)
    summary_text = getattr(summary_cell, "text", None) if summary_cell else None

    if summary_text:
        result_box.markdown("### Gap Summary")
        result_box.write(summary_text)
        st.download_button("Download Gap Report (.txt)", summary_text, file_name="gap_report.txt")
    else:
        st.info("No summary returned. Check your output column name.")